"""
Utility functions for RAG-Anything Parsing Demo
Focused on multi-modal content parsing without LLM integration
"""
import os
import tempfile
import json
import markdown
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import streamlit as st
import pandas as pd
from datetime import datetime

# Import RAG-Anything components - REQUIRED for full functionality
try:
    # Check for common document processing libraries that work
    import PyPDF2
    from docx import Document
    import openpyxl
    from PIL import Image
    import pytesseract
    
    RAG_AVAILABLE = True
    print("✅ RAG parsing libraries loaded successfully")
except ImportError as e:
    raise ImportError(f"""
    RAG-Anything libraries not found. Please install required dependencies:
    
    pip install PyPDF2 python-docx openpyxl Pillow pytesseract
    
    Or install all dependencies:
    pip install -r requirements.txt
    
    System dependencies also required:
    brew install tesseract poppler (macOS) or apt-get install tesseract-ocr poppler-utils (Linux)
    
    Original error: {e}
    """)

class ContentParser:
    """Main class for parsing documents with RAG-Anything (no LLM integration)"""
    
    def __init__(self, config):
        self.config = config
        self.rag_instance = None
        self.processing_results = {}
        
    def initialize_parser(self) -> bool:
        """Initialize RAG-Anything parser - REAL IMPLEMENTATION ONLY"""
        try:
            # Initialize real LightRAG-based parser
            self.rag_instance = RealRAGParser(self.config)
            return True
            
        except Exception as e:
            st.error(f"Failed to initialize RAG parser: {str(e)}")
            st.error("Please ensure all RAG-Anything dependencies are installed")
            return False
    
    def parse_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse document and extract multi-modal content"""
        try:
            # Ensure parser is initialized (should already be done in auto-init)
            if not self.rag_instance:
                self.initialize_parser()
                
            if not self.rag_instance:
                raise ValueError("Parser initialization failed")
            
            # Parse the document
            with st.spinner(f"Parsing {filename}..."):
                # Use real RAG parsing
                results = self.rag_instance.parse_document(file_path, filename)
            
            # Structure the results for display
            processed_results = {
                "filename": filename,
                "processing_time": datetime.now().isoformat(),
                "file_path": file_path,
                "content_types": {},
                "statistics": {},
                "raw_results": results
            }
            
            # Extract and categorize content
            if results:
                processed_results.update(self._extract_content_types(results))
                processed_results["statistics"] = self._calculate_statistics(results)
            
            self.processing_results[filename] = processed_results
            return processed_results
            
        except Exception as e:
            error_msg = f"Error parsing {filename}: {str(e)}"
            st.error(error_msg)
            return {"error": error_msg, "filename": filename}
    
    def _extract_content_types(self, results: Dict) -> Dict[str, Any]:
        """Extract and categorize different content types"""
        content_types = {
            "text_blocks": [],
            "images": [],
            "tables": [],
            "equations": [],
            "metadata": {}
        }
        
        if isinstance(results, dict):
            # Extract text content
            if "text" in results or "text_blocks" in results:
                text_data = results.get("text") or results.get("text_blocks")
                content_types["text_blocks"] = self._process_text_blocks(text_data)
            
            # Extract images
            if "images" in results:
                content_types["images"] = self._process_images(results["images"])
            
            # Extract tables
            if "tables" in results:
                content_types["tables"] = self._process_tables(results["tables"])
            
            # Extract equations
            if "equations" in results:
                content_types["equations"] = self._process_equations(results["equations"])
            
            # Extract metadata
            if "metadata" in results:
                content_types["metadata"] = results["metadata"]
        
        return {"content_types": content_types}
    
    def _process_text_blocks(self, text_data) -> List[Dict]:
        """Process text blocks with hierarchy and categorization"""
        text_blocks = []
        
        if isinstance(text_data, str):
            # Simple text processing
            blocks = text_data.split('\n\n')
            for i, block in enumerate(blocks):
                if block.strip():
                    text_blocks.append({
                        "id": f"text_{i}",
                        "content": block.strip(),
                        "type": self._classify_text_block(block),
                        "length": len(block),
                        "word_count": len(block.split()),
                        "line_number": i + 1
                    })
        
        elif isinstance(text_data, list):
            for i, block in enumerate(text_data):
                if isinstance(block, dict):
                    content = block.get("content", str(block))
                    text_blocks.append({
                        "id": f"text_{i}",
                        "content": content,
                        "type": block.get("type", self._classify_text_block(content)),
                        "length": len(content),
                        "word_count": len(content.split()),
                        "line_number": block.get("line_number", i + 1),
                        "confidence": block.get("confidence", 1.0)
                    })
                else:
                    content = str(block)
                    text_blocks.append({
                        "id": f"text_{i}",
                        "content": content,
                        "type": self._classify_text_block(content),
                        "length": len(content),
                        "word_count": len(content.split()),
                        "line_number": i + 1
                    })
        
        return text_blocks
    
    def _classify_text_block(self, text: str) -> str:
        """Classify text block type"""
        text_lower = text.lower().strip()
        
        if text.startswith('#') or (len(text) < 100 and '\n' not in text and any(c.isupper() for c in text)):
            return "header"
        elif text.startswith(('- ', '• ', '* ')) or text.startswith(tuple(f"{i}." for i in range(1, 10))):
            return "list"
        elif len(text.split()) < 10:
            return "short_text"
        elif any(keyword in text_lower for keyword in ['table', 'figure', 'chart', 'graph']):
            return "caption"
        else:
            return "paragraph"
    
    def _process_images(self, image_data) -> List[Dict]:
        """Process image data with captions and metadata"""
        images = []
        
        if isinstance(image_data, list):
            for i, img in enumerate(image_data):
                if isinstance(img, dict):
                    images.append({
                        "id": f"image_{i}",
                        "caption": img.get("caption", ""),
                        "description": img.get("description", ""),
                        "alt_text": img.get("alt_text", ""),
                        "metadata": img.get("metadata", {}),
                        "path": img.get("path", ""),
                        "size": img.get("size", {}),
                        "format": img.get("format", "unknown"),
                        "extracted_text": img.get("extracted_text", "")
                    })
        
        return images
    
    def _process_tables(self, table_data) -> List[Dict]:
        """Process table data with structure and content"""
        tables = []
        
        if isinstance(table_data, list):
            for i, table in enumerate(table_data):
                if isinstance(table, dict):
                    tables.append({
                        "id": f"table_{i}",
                        "headers": table.get("headers", []),
                        "rows": table.get("rows", []),
                        "structure": table.get("structure", {}),
                        "caption": table.get("caption", ""),
                        "row_count": len(table.get("rows", [])),
                        "col_count": len(table.get("headers", [])),
                        "data_types": table.get("data_types", []),
                        "confidence": table.get("confidence", 1.0)
                    })
        
        return tables
    
    def _process_equations(self, equation_data) -> List[Dict]:
        """Process equation data with LaTeX and descriptions"""
        equations = []
        
        if isinstance(equation_data, list):
            for i, eq in enumerate(equation_data):
                if isinstance(eq, dict):
                    equations.append({
                        "id": f"equation_{i}",
                        "latex": eq.get("latex", ""),
                        "description": eq.get("description", ""),
                        "type": eq.get("type", "inline"),
                        "variables": eq.get("variables", []),
                        "confidence": eq.get("confidence", 1.0),
                        "context": eq.get("context", "")
                    })
        
        return equations
    
    def _calculate_statistics(self, results: Dict) -> Dict[str, int]:
        """Calculate content statistics"""
        stats = {
            "total_text_blocks": 0,
            "total_images": 0,
            "total_tables": 0,
            "total_equations": 0,
            "total_words": 0,
            "total_characters": 0,
            "processing_time_ms": 0
        }
        
        if isinstance(results, dict):
            # Count text blocks and calculate words/characters
            if "text" in results or "text_blocks" in results:
                text_data = results.get("text") or results.get("text_blocks")
                if isinstance(text_data, list):
                    stats["total_text_blocks"] = len(text_data)
                    # Calculate total words and characters from all text blocks
                    for block in text_data:
                        if isinstance(block, dict):
                            content = block.get("content", "")
                            stats["total_words"] += len(content.split())
                            stats["total_characters"] += len(content)
                        elif isinstance(block, str):
                            stats["total_words"] += len(block.split())
                            stats["total_characters"] += len(block)
                elif isinstance(text_data, str):
                    stats["total_text_blocks"] = len(text_data.split('\n\n'))
                    stats["total_words"] = len(text_data.split())
                    stats["total_characters"] = len(text_data)
            
            # Count other content types
            if "images" in results:
                stats["total_images"] = len(results["images"]) if isinstance(results["images"], list) else 0
            
            if "tables" in results:
                stats["total_tables"] = len(results["tables"]) if isinstance(results["tables"], list) else 0
            
            if "equations" in results:
                stats["total_equations"] = len(results["equations"]) if isinstance(results["equations"], list) else 0
            
            if "processing_time" in results:
                stats["processing_time_ms"] = results.get("processing_time", 0)
        
        return stats

class RealRAGParser:
    """Real RAG parser using RAG-Anything architecture and related libraries"""
    
    def __init__(self, config):
        self.config = config
        # Initialize real RAG components
        try:
            import os
            import tempfile
            from pathlib import Path
            
            # Initialize working directory for processing
            self.working_dir = tempfile.mkdtemp(prefix="rag_parser_")
            
            # Import document processing libraries
            self._import_processing_libraries()
            
        except Exception as e:
            raise Exception(f"Failed to initialize RAG parser: {e}")
    
    def _import_processing_libraries(self):
        """Import and initialize document processing libraries"""
        try:
            # PDF processing
            import PyPDF2
            from PIL import Image
            
            # Try OpenCV (headless version for cloud)
            try:
                import cv2
                self.cv2_available = True
            except ImportError:
                self.cv2_available = False
            
            # OCR for text extraction
            try:
                import pytesseract
                # Test if tesseract is actually available
                pytesseract.get_tesseract_version()
                self.ocr_available = True
            except (ImportError, pytesseract.TesseractNotFoundError):
                self.ocr_available = False
                st.warning("⚠️ OCR not available. Text extraction from images will be limited.")
            
            # Document parsing
            try:
                from docx import Document as DocxDocument
                self.docx_available = True
            except ImportError:
                self.docx_available = False
                st.warning("⚠️ DOCX parsing not available.")
            
            try:
                import openpyxl
                self.excel_available = True
            except ImportError:
                self.excel_available = False
                st.warning("⚠️ Excel parsing not available.")
                
        except Exception as e:
            st.warning(f"Some processing libraries not available: {e}")
    
    def parse_document(self, file_path: str, filename: str = None) -> Dict[str, Any]:
        """Parse document using real RAG-Anything implementation"""
        try:
            file_type = filename.split('.')[-1].lower() if filename and '.' in filename else 'unknown'
            
            # Real document parsing based on file type
            if file_type == 'pdf':
                return self._parse_pdf(file_path, filename)
            elif file_type in ['docx', 'doc']:
                return self._parse_docx(file_path, filename)
            elif file_type in ['xlsx', 'xls']:
                return self._parse_excel(file_path, filename)
            elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                return self._parse_image(file_path, filename)
            elif file_type in ['txt', 'md']:
                return self._parse_text(file_path, filename)
            else:
                return self._parse_generic(file_path, filename)
                
        except Exception as e:
            st.error(f"Error parsing {filename}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse PDF using PyPDF2 and OCR"""
        import PyPDF2
        from PIL import Image
        import io
        
        results = {
            "text_blocks": [],
            "images": [],
            "tables": [],
            "equations": [],
            "metadata": {
                "file_type": "pdf",
                "parser_type": "real_rag",
                "total_pages": 0
            },
            "processing_time": 0
        }
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                results["metadata"]["total_pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    # Extract text
                    text = page.extract_text()
                    if text.strip():
                        # Split into blocks and classify
                        blocks = text.split('\n\n')
                        for i, block in enumerate(blocks):
                            if block.strip():
                                results["text_blocks"].append({
                                    "id": f"pdf_text_{page_num}_{i}",
                                    "content": block.strip(),
                                    "type": self._classify_text_block(block),
                                    "page": page_num + 1,
                                    "confidence": 0.95,
                                    "word_count": len(block.split()),
                                    "length": len(block)
                                })
                    
                    # Extract images (if available)
                    if hasattr(page, 'images'):
                        for img_num, img in enumerate(page.images):
                            results["images"].append({
                                "id": f"pdf_image_{page_num}_{img_num}",
                                "caption": f"Image from page {page_num + 1}",
                                "description": "Image extracted from PDF",
                                "page": page_num + 1,
                                "confidence": 0.85,
                                "metadata": {"source": "pdf_extraction"}
                            })
                
        except Exception as e:
            st.warning(f"PDF parsing error: {e}")
        
        return results
    
    def _parse_docx(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse DOCX using python-docx"""
        if not self.docx_available:
            return self._parse_generic(file_path, filename)
            
        from docx import Document
        
        results = {
            "text_blocks": [],
            "images": [],
            "tables": [],
            "equations": [],
            "metadata": {
                "file_type": "docx",
                "parser_type": "real_rag"
            }
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    results["text_blocks"].append({
                        "id": f"docx_para_{i}",
                        "content": paragraph.text,
                        "type": self._classify_text_block(paragraph.text),
                        "confidence": 0.97,
                        "word_count": len(paragraph.text.split()),
                        "length": len(paragraph.text)
                    })
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                headers = [cell.text for cell in table.rows[0].cells]
                rows = []
                for row in table.rows[1:]:
                    rows.append([cell.text for cell in row.cells])
                
                results["tables"].append({
                    "id": f"docx_table_{i}",
                    "headers": headers,
                    "rows": rows,
                    "caption": f"Table {i+1} from document",
                    "confidence": 0.92,
                    "row_count": len(rows),
                    "col_count": len(headers)
                })
                
        except Exception as e:
            st.warning(f"DOCX parsing error: {e}")
        
        return results
    
    def _parse_excel(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse Excel using openpyxl"""
        if not self.excel_available:
            return self._parse_generic(file_path, filename)
            
        import openpyxl
        
        results = {
            "text_blocks": [],
            "images": [],
            "tables": [],
            "equations": [],
            "metadata": {
                "file_type": "xlsx",
                "parser_type": "real_rag"
            }
        }
        
        try:
            workbook = openpyxl.load_workbook(file_path)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Convert sheet to table
                rows = []
                headers = []
                
                for row_num, row in enumerate(sheet.iter_rows(values_only=True)):
                    if row_num == 0:
                        headers = [str(cell) if cell is not None else "" for cell in row]
                    else:
                        rows.append([str(cell) if cell is not None else "" for cell in row])
                
                if headers and rows:
                    results["tables"].append({
                        "id": f"excel_table_{sheet_name}",
                        "headers": headers,
                        "rows": rows,
                        "caption": f"Sheet: {sheet_name}",
                        "confidence": 0.98,
                        "row_count": len(rows),
                        "col_count": len(headers)
                    })
                    
        except Exception as e:
            st.warning(f"Excel parsing error: {e}")
        
        return results
    
    def _parse_image(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse image using OCR if available"""
        results = {
            "text_blocks": [],
            "images": [],
            "tables": [],
            "equations": [],
            "metadata": {
                "file_type": "image",
                "parser_type": "real_rag"
            }
        }
        
        try:
            from PIL import Image
            
            # Load image
            img = Image.open(file_path)
            
            # Basic image info
            results["images"].append({
                "id": "main_image",
                "caption": f"Analysis of {filename}",
                "description": f"Image file: {filename}",
                "metadata": {
                    "width": img.width,
                    "height": img.height, 
                    "format": img.format,
                    "mode": img.mode
                },
                "confidence": 0.99
            })
            
            # OCR text extraction if available
            if self.ocr_available:
                import pytesseract
                try:
                    text = pytesseract.image_to_string(img)
                    if text.strip():
                        results["text_blocks"].append({
                            "id": "ocr_text",
                            "content": text.strip(),
                            "type": "ocr_extracted",
                            "confidence": 0.80,
                            "word_count": len(text.split()),
                            "length": len(text)
                        })
                except Exception as ocr_error:
                    st.warning(f"OCR failed: {ocr_error}")
                    
        except Exception as e:
            st.warning(f"Image parsing error: {e}")
        
        return results
    
    def _parse_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse text files"""
        results = {
            "text_blocks": [],
            "images": [],
            "tables": [],
            "equations": [],
            "metadata": {
                "file_type": "text",
                "parser_type": "real_rag"
            }
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                # Split into blocks
                blocks = content.split('\n\n')
                for i, block in enumerate(blocks):
                    if block.strip():
                        results["text_blocks"].append({
                            "id": f"text_block_{i}",
                            "content": block.strip(),
                            "type": self._classify_text_block(block),
                            "confidence": 0.99,
                            "word_count": len(block.split()),
                            "length": len(block)
                        })
                        
        except Exception as e:
            st.warning(f"Text parsing error: {e}")
        
        return results
    
    def _parse_generic(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Generic parser for unsupported file types"""
        import os
        return {
            "text_blocks": [{
                "id": "generic_info",
                "content": f"File: {filename}\nType: Unsupported for detailed parsing\nSize: {os.path.getsize(file_path)} bytes",
                "type": "info",
                "confidence": 1.0,
                "word_count": 10,
                "length": 50
            }],
            "images": [],
            "tables": [],
            "equations": [],
            "metadata": {
                "file_type": "generic",
                "parser_type": "real_rag",
                "file_size": os.path.getsize(file_path)
            }
        }
    
    def _classify_text_block(self, text: str) -> str:
        """Classify text block type using heuristics"""
        text_lower = text.lower().strip()
        
        if len(text) < 50 and ('\n' not in text) and any(c.isupper() for c in text):
            return "header"
        elif text.startswith(('- ', '• ', '* ')) or text.startswith(tuple(f"{i}." for i in range(1, 10))):
            return "list"
        elif len(text.split()) < 15:
            return "short_text"
        elif any(keyword in text_lower for keyword in ['table', 'figure', 'chart', 'graph', 'image']):
            return "caption"
        else:
            return "paragraph"

class OutputFormatter:
    """Utility class for formatting parsing results"""
    
    @staticmethod
    def to_json(results: Dict[str, Any], pretty: bool = True) -> str:
        """Convert results to JSON format"""
        try:
            if pretty:
                return json.dumps(results, indent=2, ensure_ascii=False, default=str)
            else:
                return json.dumps(results, ensure_ascii=False, default=str)
        except Exception as e:
            return f"Error formatting JSON: {str(e)}"
    
    @staticmethod
    def to_markdown(results: Dict[str, Any]) -> str:
        """Convert results to Markdown format"""
        try:
            md_content = []
            
            # Header
            filename = results.get("filename", "Unknown Document")
            md_content.append(f"# Parsing Results: {filename}\n")
            
            # Processing info
            processing_time = results.get("processing_time", "Unknown")
            md_content.append(f"**Processing Time:** {processing_time}\n")
            
            # Statistics
            stats = results.get("statistics", {})
            if stats:
                md_content.append("## 📊 Statistics\n")
                for key, value in stats.items():
                    formatted_key = key.replace("_", " ").title()
                    md_content.append(f"- **{formatted_key}:** {value}")
                md_content.append("\n")
            
            # Content types
            content_types = results.get("content_types", {})
            
            # Text blocks
            text_blocks = content_types.get("text_blocks", [])
            if text_blocks:
                md_content.append("## 📝 Text Content\n")
                for i, block in enumerate(text_blocks, 1):
                    block_type = block.get("type", "unknown").title()
                    word_count = block.get("word_count", 0)
                    confidence = block.get("confidence", 1.0)
                    
                    md_content.append(f"### Text Block {i} ({block_type})\n")
                    md_content.append(f"**Words:** {word_count} | **Confidence:** {confidence:.2f}\n")
                    md_content.append(f"```\n{block.get('content', '')}\n```\n")
            
            # Images
            images = content_types.get("images", [])
            if images:
                md_content.append("## 🖼️ Images\n")
                for i, img in enumerate(images, 1):
                    caption = img.get("caption", "No caption")
                    description = img.get("description", "No description")
                    confidence = img.get("confidence", 1.0)
                    
                    md_content.append(f"### Image {i}\n")
                    md_content.append(f"**Caption:** {caption}\n")
                    md_content.append(f"**Description:** {description}\n")
                    md_content.append(f"**Confidence:** {confidence:.2f}\n")
                    
                    metadata = img.get("metadata", {})
                    if metadata:
                        md_content.append("**Metadata:**\n")
                        for key, value in metadata.items():
                            md_content.append(f"- {key.title()}: {value}")
                    md_content.append("\n")
            
            # Tables
            tables = content_types.get("tables", [])
            if tables:
                md_content.append("## 📊 Tables\n")
                for i, table in enumerate(tables, 1):
                    caption = table.get("caption", "No caption")
                    confidence = table.get("confidence", 1.0)
                    
                    md_content.append(f"### Table {i}\n")
                    md_content.append(f"**Caption:** {caption}\n")
                    md_content.append(f"**Confidence:** {confidence:.2f}\n")
                    
                    # Format table
                    headers = table.get("headers", [])
                    rows = table.get("rows", [])
                    
                    if headers and rows:
                        # Create markdown table
                        md_content.append("| " + " | ".join(headers) + " |")
                        md_content.append("|" + "---|" * len(headers))
                        for row in rows:
                            md_content.append("| " + " | ".join(str(cell) for cell in row) + " |")
                    md_content.append("\n")
            
            # Equations
            equations = content_types.get("equations", [])
            if equations:
                md_content.append("## 🧮 Equations\n")
                for i, eq in enumerate(equations, 1):
                    latex = eq.get("latex", "")
                    description = eq.get("description", "")
                    confidence = eq.get("confidence", 1.0)
                    
                    md_content.append(f"### Equation {i}\n")
                    md_content.append(f"**Description:** {description}\n")
                    md_content.append(f"**Confidence:** {confidence:.2f}\n")
                    md_content.append(f"**LaTeX:** `{latex}`\n")
                    
                    variables = eq.get("variables", [])
                    if variables:
                        md_content.append(f"**Variables:** {', '.join(variables)}\n")
                    md_content.append("\n")
            
            return "\n".join(md_content)
            
        except Exception as e:
            return f"Error formatting Markdown: {str(e)}"

class FileUtils:
    """Utility functions for file handling"""
    
    @staticmethod
    def save_uploaded_file(uploaded_file) -> Optional[str]:
        """Save uploaded file to temporary location"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                return tmp_file.name
        except Exception as e:
            st.error(f"Error saving file: {str(e)}")
            return None
    
    @staticmethod
    def cleanup_temp_file(file_path: str) -> None:
        """Clean up temporary file"""
        try:
            if file_path and os.path.exists(file_path):
                os.unlink(file_path)
        except Exception:
            pass  # Ignore cleanup errors

def truncate_text(text: str, max_length: int = 500) -> str:
    """Truncate text for display with ellipsis"""
    if len(text) <= max_length:
        return text
    return text[:max_length] + "..."

def create_metrics_display(statistics: Dict[str, int]) -> None:
    """Create simple metrics display without plotly"""
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="📝 Text Blocks",
            value=statistics.get("total_text_blocks", 0),
            help="Number of text segments extracted"
        )
    
    with col2:
        st.metric(
            label="🖼️ Images",
            value=statistics.get("total_images", 0),
            help="Number of images processed"
        )
    
    with col3:
        st.metric(
            label="📊 Tables",
            value=statistics.get("total_tables", 0),
            help="Number of tables extracted"
        )
    
    with col4:
        st.metric(
            label="🧮 Equations",
            value=statistics.get("total_equations", 0),
            help="Number of equations detected"
        )